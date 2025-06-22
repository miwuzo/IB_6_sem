from docx import Document

def read_docx(file_path):
    """Читает текст из .docx файла."""
    doc = Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs])

def write_docx(file_path, text):
    """Записывает текст в .docx файл."""
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_path)

def encode_message(container_file, message_file, stego_file, method='variable'):
    """Кодирует сообщение с использованием выбранного метода."""
    try:
        container_text = read_docx(container_file)
        message = read_docx(message_file)
    except FileNotFoundError as e:
        print(f"Ошибка: Файл не найден: {e}")
        return
    except Exception as e:
        print(f"Ошибка при чтении файлов: {e}")
        return

    # Подготовка сообщения: Перевод в биты
    binary_message = ''.join(format(ord(char), '08b') for char in message)

    # Разбиение контейнера на слова
    words = container_text.split()
    stego_words = []
    bit_index = 0

    for word in words:
        if method == 'variable':
            # Метод переменной длины слов
            stego_words.append(word)
            if bit_index < len(binary_message):
                if len(word) < 5:  # Короткое слово
                    bit = binary_message[bit_index]
                    stego_words.append("  " if bit == '1' else " ")
                    bit_index += 1
                else:  # Длинное слово
                    if bit_index + 1 < len(binary_message):
                        bit1 = binary_message[bit_index]
                        bit2 = binary_message[bit_index + 1]
                        if bit1 == '0' and bit2 == '0':
                            stego_words.append(" ")
                        elif bit1 == '0' and bit2 == '1':
                            stego_words.append("   ")
                        elif bit1 == '1' and bit2 == '0':
                            stego_words.append("    ")
                        else:
                            stego_words.append("     ")
                        bit_index += 2
        elif method == 'aprosha':
            # Метод апроша
            new_word = ""
            for char in word:
                new_word += char
                if bit_index < len(binary_message):
                    if binary_message[bit_index] == '1':
                        new_word += " "  # Увеличиваем расстояние
                    bit_index += 1
            stego_words.append(new_word)

    stego_text = " ".join(stego_words)

    try:
        write_docx(stego_file, stego_text)
        print(f"Сообщение успешно закодировано и сохранено в {stego_file}")
    except Exception as e:
        print(f"Ошибка при записи в файл: {e}")

def decode_message(stego_file, method='variable'):
    """Извлекает скрытое сообщение из стеготекста с использованием выбранного метода."""
    try:
        stego_text = read_docx(stego_file)
    except FileNotFoundError as e:
        print(f"Ошибка: Файл не найден: {e}")
        return None
    except Exception as e:
        print(f"Ошибка при чтении файла: {e}")
        return None

    binary_message = ""
    words = stego_text.split()

    for word in words:
        if method == 'variable':
            if len(word) < 5:  # Короткое слово
                if words[words.index(word) + 1] == "":  # Два пробела
                    binary_message += '1'
                elif words[words.index(word) + 1] == " ":  # Один пробел
                    binary_message += '0'
            else:  # Длинное слово
                next_word = words[words.index(word) + 1]
                if next_word == " ":
                    binary_message += '00'
                elif next_word == "   ":
                    binary_message += '01'
                elif next_word == "    ":
                    binary_message += '10'
                elif next_word == "     ":
                    binary_message += '11'

        elif method == 'aprosha':
            for char in word:
                if char == ' ':
                    binary_message += '1'  # Увеличенное расстояние
                else:
                    binary_message += '0'  # Обычное расстояние

    # Преобразование битов в символы
    decoded_message = ""
    for i in range(0, len(binary_message), 8):
        byte = binary_message[i:i + 8]
        if len(byte) == 8:
            try:
                char_code = int(byte, 2)
                decoded_message += chr(char_code)
            except ValueError:
                print(f"Ошибка: Некорректная битовая последовательность: {byte}")
                return None

    return decoded_message

# Пример использования:
if __name__ == "__main__":
    container_file = "container.docx"
    message_file = "message.docx"
    stego_file = "stego.docx"
    stego_file2 = "stego2.docx"
    write_docx(container_file, "Пример текста для контейнера.")
    write_docx(message_file, "секрет")

    # Кодирование с использованием метода переменной длины
    encode_message(container_file, message_file, stego_file, method='variable')


    # Кодирование с использованием метода апроша
    encode_message(container_file, message_file, stego_file2, method='aprosha')
